from datetime import datetime, date, timedelta
import re

import Users
from docx import Document
import os
import requests

googl_key = os.getenv("GOOGL_KEY")



class Translator:
    def transl (word: str):
        url = f"https://translation.googleapis.com/language/translate/v2?key={googl_key}"
        if re.search(r"[а-яА-ЯёЁ]", word):
            data = {"q": word, "target": "de"}
            respinse = requests.post(url, data=data)
            result = respinse.json()
            return result["data"]["translations"][0]["translatedText"]


        elif re.search(r"[a-zA-ZäöüÄÖÜß]", word):
            data = {"q": word, "target": "ru"}
            respinse = requests.post(url, data=data)
            result = respinse.json()
            return result["data"]["translations"][0]["translatedText"]

    #def transl(word: str):
    #    auth_key = "677b176c-d12c-432d-8186-85674df5d49d:fx"
     #   deepl_client = deepl.DeepLClient(auth_key)
     #   if re.search(r"[а-яА-ЯёЁ]", word):

     #       result = deepl_client.translate_text(word, source_lang="Ru", target_lang="DE")
     #       return result.text
     #   elif re.search(r"[a-zA-ZäöüÄÖÜß]", word):
      #      result = deepl_client.translate_text(word, source_lang="DE", target_lang="Ru")
       #     return result.text

    def show_vocabular(user: Users.User):
        if not user.vocabular:
            return "Словарь пуст."

        result = []
        for word, data in user.vocabular.items():
            meaning = data.get("meaning", "—")
            result.append(f"{word} — {meaning}")
        k = len(result)
        return f'Вего слов : {k}\n\n' + "\n".join(result)

    def show_todays_vocabular(user: Users.User):
        if not user.vocabular:
            return "Словарь пуст."

        today = date.today()
        result = []

        for word, data in user.vocabular.items():
            added_at = data.get("added_at")
            if not added_at:
                continue
            added_date = None
            if added_date is None:
                try:
                    added_date = datetime.strptime(added_at, "%Y-%m-%d %H:%M:%S").date()
                except Exception:
                    continue

            if added_date == today:
                meaning = data.get("meaning", "—")
                result.append(f"{word} — {meaning}")

        k = len(result)
        if k == 0:
            return "Сегодня слов нет."
        return f"Всего слов сегодня: {k}\n\n" + "\n".join(result)



    @staticmethod
    def export_vocabular_to_docx(user):
        """Создаёт DOCX-файл со словами пользователя"""
        if not user.vocabular:
            return None

        doc = Document()
        doc.add_heading(f"Словарь пользователя {user.name}", level=1)
        doc.add_paragraph(f"Экспорт от {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "№"
        hdr_cells[1].text = "Слово"
        hdr_cells[2].text = "Перевод"

        for i, (word, data) in enumerate(user.vocabular.items(), start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = word
            row_cells[2].text = data.get("meaning", "—")

        # имя файла — уникальное по user_id
        filename = f"user_{user.user_id}_vocabular.docx"
        filepath = os.path.join(os.getcwd(), filename)
        doc.save(filepath)

        return filepath
